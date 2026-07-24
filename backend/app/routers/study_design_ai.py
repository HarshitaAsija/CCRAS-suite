from fastapi import APIRouter, HTTPException, Query, status, Depends
from pydantic import BaseModel
from typing import Dict, Any, List, Optional
from datetime import datetime
from sqlalchemy.orm import Session
from app.db.session import get_db
from app.services.study_design_ai import ai_service

router = APIRouter(prefix="/study-design", tags=["study-design"])

# --- INPUT SCHEMAS ---
class PicoRequest(BaseModel):
    population: str
    intervention: str
    comparator: Optional[str] = ""
    outcome: str

class StatPlanRequest(BaseModel):
    studyType: str
    outcome: str

class CriteriaRequest(BaseModel):
    population: str
    intervention: str

class ConfoundersRequest(BaseModel):
    population: str
    intervention: str

class AyurvedaRequest(BaseModel):
    intervention: str

class TimelineRequest(BaseModel):
    studyType: str
    durationWeeks: int = 12
    researchQuestion: Optional[str] = None
    pico: Optional[Dict[str, Any]] = None

class ExportRequest(BaseModel):
    state: Dict[str, Any]
    format: str  # "markdown", "json", "html"

# --- ENDPOINTS ---

@router.post("/analyze")
def analyze_protocol_endpoint(payload: Dict[str, Any]):
    try:
        result = ai_service.analyze_protocol(payload)
        # Format dates / temporal fields manually to avoid SQLAlchemy issues
        result["lastCalculated"] = datetime.now().isoformat()
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

@router.post("/generate-hypothesis")
def generate_hypothesis_endpoint(payload: PicoRequest, db: Session = Depends(get_db)):
    try:
        return ai_service.generate_hypothesis(payload.model_dump(), db)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/recommend-study-type")
def recommend_study_type_endpoint(payload: PicoRequest, db: Session = Depends(get_db)):
    try:
        return ai_service.recommend_study_type(payload.model_dump(), db)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/statistical-plan")
def recommend_stat_plan_endpoint(payload: StatPlanRequest, db: Session = Depends(get_db)):
    try:
        return ai_service.recommend_statistical_plan(payload.studyType, payload.outcome, db)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/suggest-criteria")
def suggest_criteria_endpoint(payload: CriteriaRequest, db: Session = Depends(get_db)):
    try:
        return ai_service.suggest_criteria(payload.model_dump(), db)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/suggest-confounders")
def suggest_confounders_endpoint(payload: ConfoundersRequest, db: Session = Depends(get_db)):
    try:
        return ai_service.suggest_confounders(payload.model_dump(), db)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/suggest-ayurveda")
def suggest_ayurveda_endpoint(payload: AyurvedaRequest, db: Session = Depends(get_db)):
    try:
        return ai_service.suggest_ayurveda_protocol(payload.intervention, db)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/suggest-timeline")
def suggest_timeline_endpoint(payload: TimelineRequest, db: Session = Depends(get_db)):
    try:
        return ai_service.suggest_timeline(
            payload.studyType, 
            payload.durationWeeks, 
            payload.researchQuestion, 
            payload.pico, 
            db
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/bias-risk")
def assess_bias_risk_endpoint(payload: Dict[str, Any]):
    try:
        return ai_service.assess_bias_risk(payload)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/guideline-check")
def guideline_check_endpoint(payload: Dict[str, Any]):
    try:
        return ai_service.check_guidelines(payload)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/protocol-summary")
def protocol_summary_endpoint(payload: Dict[str, Any]):
    try:
        return ai_service.generate_detailed_protocol_summary(payload)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/export")
def export_protocol_endpoint(payload: ExportRequest):
    try:
        state = payload.state
        fmt = payload.format.lower()
        
        # Use detailed protocol summary
        detailed_summary = ai_service.generate_detailed_protocol_summary(state)
        
        if fmt == "markdown":
            md = f"# Study Protocol Export\n\n"
            md += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            
            for section in detailed_summary['sections']:
                md += f"## {section['title']}\n\n{section['content']}\n\n"
            
            return {
                "format": "markdown",
                "content": md,
                "filename": f"protocol_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
            }
        
        elif fmt == "json":
            return {
                "format": "json",
                "content": detailed_summary,
                "filename": f"protocol_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            }
        
        elif fmt == "html":
            html = f"<html><head><title>Study Protocol Export</title></head><body style='font-family: sans-serif; padding: 20px; max-width: 800px; margin: auto;'>"
            html += f"<h1>Study Protocol Export</h1><p><em>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</em></p><hr/>"
            
            for section in detailed_summary['sections']:
                # Basic markdown to html conversion for the simple formatting used
                content = section['content'].replace('\n\n', '</p><p>').replace('\n', '<br/>')
                # Fix bold tags
                import re
                content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
                html += f"<h2>{section['title']}</h2><p>{content}</p>"
                
            html += "</body></html>"
            return {
                "format": "html",
                "content": html,
                "filename": f"protocol_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            }
            
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


@router.post("/export-docx")
def export_docx_endpoint(payload: Dict[str, Any]):
    try:
        from docx import Document
        import io
        from fastapi import Response
        
        state = payload.get("state", payload)
        
        doc = Document()
        doc.add_heading(state.get('title', 'Untitled Study Protocol'), level=0)
        
        # 1. Background
        doc.add_heading("1. Background & Research Question", level=1)
        doc.add_paragraph(state.get('research_question') or 'Not specified')
        
        # 2. PICO Framework
        pico = state.get('pico', {})
        doc.add_heading("2. PICO Framework", level=1)
        doc.add_paragraph(f"Population: {pico.get('population', 'Not specified')}")
        doc.add_paragraph(f"Intervention: {pico.get('intervention', 'Not specified')}")
        doc.add_paragraph(f"Comparator: {pico.get('comparator', 'Not specified')}")
        doc.add_paragraph(f"Outcome: {pico.get('outcome', 'Not specified')}")
        
        # 3. Hypotheses
        hyp = state.get('hypothesis', {})
        doc.add_heading("3. Objectives & Hypotheses", level=1)
        doc.add_paragraph(f"Primary Objective: {hyp.get('primaryObjective', 'Not specified')}")
        doc.add_paragraph(f"Primary Hypothesis (H1): {hyp.get('primary', 'Not specified')}")
        doc.add_paragraph(f"Null Hypothesis (H0): {hyp.get('nullHypothesis', 'Not specified')}")
        doc.add_paragraph(f"Alternative Hypothesis: {hyp.get('alternative', 'Not specified')}")
        
        # 4. Study Design
        doc.add_heading("4. Study Design", level=1)
        doc.add_paragraph(f"Recommended Design: {state.get('study_type', {}).get('recommended', 'Not specified')}")
        
        # 5. Sample Size
        sample_params = state.get('sample_size', state.get('sampleSizeParams', {}))
        sample_result = state.get('sample_size_result', state.get('sampleSizeResult', {}))
        doc.add_heading("5. Sample Size Calculation", level=1)
        doc.add_paragraph(f"Alpha (Type I error): {sample_params.get('alpha', 0.05)}")
        doc.add_paragraph(f"Power (1 - Type II error): {sample_params.get('power', 0.8)}")
        doc.add_paragraph(f"Effect Size (Cohen's d): {sample_params.get('effectSize', 0.5)}")
        doc.add_paragraph(f"Calculated Total Sample Size: {sample_result.get('total', 'Not calculated')}")
        doc.add_paragraph(f"Per Arm Sample Size: {sample_result.get('perArm', 'Not calculated')}")
        
        # 6. Statistical Analysis Plan
        stat = state.get('statistical_plan', state.get('statisticalPlan', {}))
        doc.add_heading("6. Statistical Analysis Plan", level=1)
        doc.add_paragraph(f"Primary Endpoint: {stat.get('primaryEndpoint', 'Not specified')}")
        doc.add_paragraph(f"Recommended Statistical Test: {stat.get('recommendedTest', 'Not specified')}")
        doc.add_paragraph(f"Missing Data Handling: {stat.get('missingData', 'Not specified')}")
        doc.add_paragraph(f"Regression Adjustments: {stat.get('regression', 'Not specified')}")
        
        # 7. Eligibility Criteria
        crit = state.get('eligibility', state.get('criteria', {}))
        doc.add_heading("7. Eligibility Criteria", level=1)
        doc.add_heading("Inclusion Criteria", level=2)
        for c in crit.get('inclusion', []):
            doc.add_paragraph(f"- {c.get('text', c)}")
        doc.add_heading("Exclusion Criteria", level=2)
        for c in crit.get('exclusion', []):
            doc.add_paragraph(f"- {c.get('text', c)}")
            
        # 8. Potential Confounders
        doc.add_heading("8. Confounders & Mitigations", level=1)
        for c in state.get('confounders', []):
            doc.add_paragraph(f"- {c.get('name', 'Confounder')}: {c.get('mitigation', '')} (Risk level: {c.get('risk', '')})")
            
        # 9. Ayurveda Protocol
        ayur = state.get('ayush_protocol', state.get('ayurveda', {}))
        if ayur.get('formulation'):
            doc.add_heading("9. Ayurveda Protocol Specifications", level=1)
            doc.add_paragraph(f"Formulation: {ayur.get('formulation', 'Not specified')}")
            doc.add_paragraph(f"Dosage & Administration: {ayur.get('dosage', 'Not specified')}")
            doc.add_paragraph(f"Anupana (Carrier): {ayur.get('anupana', 'Not specified')}")
            doc.add_paragraph(f"Prakriti Suitability: {ayur.get('prakriti', 'Not specified')}")
            doc.add_paragraph(f"Standardization compliance: {ayur.get('standardization', 'Not specified')}")
            doc.add_paragraph(f"Safety/Toxicity monitors: {ayur.get('safety', 'Not specified')}")
            
        # 10. Study Timeline
        doc.add_heading("10. Timeline & Milestones", level=1)
        for t in state.get('timeline', []):
            doc.add_paragraph(f"- {t.get('label')}: {t.get('duration')}")
            
        # 11. Ethical Milestones
        doc.add_heading("11. Ethics & Registries Milestones", level=1)
        for e in state.get('ethics', []):
            status_text = "[X]" if e.get('checked') else "[ ]"
            doc.add_paragraph(f"{status_text} {e.get('label')}")
            
        # 12. Quality Audit
        intel = state.get('intelligence', {})
        doc.add_heading("12. Protocol Audit Results", level=1)
        doc.add_paragraph(f"Quality Score: {intel.get('qualityScore', 0)} / 100")
        doc.add_paragraph(f"Completeness Rate: {intel.get('completeness', 0)}%")
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": "attachment; filename=protocol_report.docx",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Docx compilation failed: {str(e)}")
